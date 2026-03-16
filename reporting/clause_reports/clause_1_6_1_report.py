"""
reporting/clause_1_6_1/report.py

Full ITSAR 1.6.1 report generator.
All static document text is loaded from report_content.json (co-located with this file).
"""

import os
import json
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from reporting.front_page import add_front_page


# ---------------------------------------------------------------------------
# Content loader
# ---------------------------------------------------------------------------

_CONTENT_JSON = os.path.abspath(os.path.join(
    os.path.dirname(__file__),
    "..",
    "report_content",
    "clause_1_6_1.json"
))

def _load_content() -> dict:
    with open(_CONTENT_JSON, "r", encoding="utf-8") as fh:
        return json.load(fh)


# ---------------------------------------------------------------------------
# Low-level format helpers  (previously scattered across master.py)
# ---------------------------------------------------------------------------

def _add_itsar_heading(doc, text: str, level: int):
    """Bold purple heading with a purple underline border."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.color.rgb = RGBColor(75, 0, 130)

    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(2)

    p = para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "4B0082")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_itsar_subheading(doc, text: str, level: int):
    """Bold purple subheading without underline border."""
    para = doc.add_heading(text, level=level)
    run = para.runs[0]
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.color.rgb = RGBColor(75, 0, 130)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(8)
    return para


def _keep_with_next(para):
    para.paragraph_format.keep_with_next = True
    para.paragraph_format.keep_together = True
    return para


def _add_bold_paragraph(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def _style_table_header(cell, bg_color: str = "1F4E79"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), bg_color)
    tc_pr.append(shd)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)


def _prevent_table_row_split(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit")
        trPr.append(cantSplit)


def _normalize_list(items) -> list:
    if not items:
        return ["None"]
    cleaned = [i.strip() for i in items if i.strip()]
    return cleaned if cleaned else ["None"]


def _add_screenshot_evidence_block(doc, title: str, image_path: str):
    TABLE_WIDTH = Inches(7.8)
    IMAGE_WIDTH = Inches(6.2)

    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    _prevent_table_row_split(table)

    table.columns[0].width = TABLE_WIDTH
    for row in table.rows:
        row.cells[0].width = TABLE_WIDTH

    for row in table.rows:
        cell = row.cells[0]
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F3ECFA")
        tcPr.append(shd)
        cell.top_margin = Inches(0.2)
        cell.bottom_margin = Inches(0.2)
        cell.left_margin = Inches(0.3)
        cell.right_margin = Inches(0.3)

    # Title row
    title_cell = table.cell(0, 0)
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 0, 130)

    # Image row
    img_cell = table.cell(1, 0)
    p_img = img_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.keep_together = True
    p_img.add_run().add_picture(image_path, width=IMAGE_WIDTH)

    # Purple border
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:color"), "4B0082")
        tblBorders.append(border)
    tblPr.append(tblBorders)
    return table


def _add_two_col_cipher_table(doc, strong_items, weak_items,
                               strong_header: str = "Strong",
                               weak_header: str = "Weak"):
    """Generic 2-column strong/weak cipher table."""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = strong_header
    t.cell(0, 1).text = weak_header
    row = t.add_row().cells
    row[0].text = "\n".join(_normalize_list(strong_items))
    row[1].text = "\n".join(_normalize_list(weak_items))
    _prevent_table_row_split(t)
    return t


def _add_result_table_header(rt, headers, bg_color: str = "4B0082"):
    for i, header in enumerate(headers):
        cell = rt.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _style_table_header(cell, bg_color=bg_color)
        cell.top_margin = Inches(0.15)
        cell.bottom_margin = Inches(0.15)
        cell.left_margin = Inches(0.15)
        cell.right_margin = Inches(0.15)


def _pad_data_rows(rt):
    for row in rt.rows[1:]:
        for cell in row.cells:
            cell.top_margin = Inches(0.12)
            cell.bottom_margin = Inches(0.12)
            cell.left_margin = Inches(0.12)
            cell.right_margin = Inches(0.12)


# ---------------------------------------------------------------------------
# Report class
# ---------------------------------------------------------------------------

class Clause161Report:
    """
    Generates the ITSAR 1.6.1 Cryptographic Based Secure Communication report.

    Parameters
    ----------
    context : object
        Must expose:
            context.evidence.run_dir  – output directory
            context.start_time        – datetime or formatted string
            context.end_time          – datetime or formatted string
    results : dict
        Must contain keys:
            dut_info         – dict with dut_name, dut_version, os_hash, config_hash
            nmap             – run_nmap_scan() result
            cipher           – run_cipher_detection() result
            ssh              – run_ssh_verification() result
            weak_cipher      – run_ssh_weak_cipher_test() result
            https_cipher     – run_httpsCipher_detection() result
            https            – run_tls_verification() result
    """

    def __init__(self, context, results: dict):
        self.context = context
        self.results = results
        print(results)
        self.output_dir = context.evidence.run_dir
        self.content = _load_content()

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def generate(self) -> str:
        """Build the document and return the path to the saved .docx file."""
        c = self.content
        r = self.context.scan_results

        doc = Document()
        dut_info = r["dut_info"]

        # Determine overall pass/fail from SSH + HTTPS results
        tc1_ssh   = r["cipher"].get("result", "FAIL")
        tc2_ssh   = r["ssh"].get("final_result", "FAIL")
        tc3_ssh   = self._compute_weak_cipher_result(r["weak_cipher"])
        tc1_https = r["https_cipher"].get("result", "FAIL")
        tc2_https = r["https"].get("final_result", "FAIL")

        final_result = (
            "PASS"
            if all(x == "PASS" for x in [tc1_ssh, tc2_ssh, tc3_ssh, tc1_https, tc2_https])
            else "FAIL"
        )

        # ---- Front page ----
        add_front_page(doc, meta={
            "dut_name":       dut_info["dut_name"],
            "dut_version":    dut_info["dut_version"],
            "os_hash":        dut_info["os_hash"],
            "config_hash":    dut_info["config_hash"],
            "start_time":     str(self.context.start_time),
            "end_time":       "",
            "final_result":   final_result,
            "itsar_id":       c["meta"]["itsar_id"],
            "itsar_version":  c["meta"]["itsar_version"],
        })

        # ---- Sections 1–8 ----
        self._add_intro_sections(doc)

        # ---- Section 9: SSH ----
        self._add_ssh_section(doc, r, tc1_ssh, tc2_ssh, tc3_ssh)

        # ---- Section 10–11: SSH observation + result table ----
        self._add_ssh_observation(doc, tc1_ssh, tc2_ssh, tc3_ssh)
        self._add_ssh_result_table(doc, tc1_ssh, tc2_ssh, tc3_ssh, r["weak_cipher"])

        # ---- Section 12: HTTPS ----
        self._add_https_section(doc, r, tc1_https, tc2_https)

        # ---- Section 13–14: HTTPS observation + result table ----
        self._add_https_observation(doc, tc1_https, tc2_https)
        self._add_https_result_table(doc, tc1_https, tc2_https)

        # ---- Save ----
        filename = (
            "ITSAR_2.6.1_Cryptographic_Based_Secure_Communication_Report_"
            f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        report_path = os.path.join(self.output_dir, filename)
        doc.save(report_path)
        return report_path

    # ------------------------------------------------------------------
    # Sections 1–8  (intro / context)
    # ------------------------------------------------------------------

    def _add_intro_sections(self, doc):
        s = self.content["sections"]

        # Section 1
        _add_itsar_heading(doc, s["1_access_authorization"]["heading"], 2)
        doc.add_paragraph(s["1_access_authorization"]["body"])

        # Section 2
        _add_itsar_heading(doc, s["2_cryptographic_communication"]["heading"], 2)
        doc.add_paragraph(s["2_cryptographic_communication"]["body"])

        # Section 3
        _add_itsar_heading(doc, s["3_requirement_description"]["heading"], 2)
        doc.add_paragraph(s["3_requirement_description"]["body"])

        # Section 4
        nmap_data = self.context.scan_results["nmap"]
        cfg = s["4_dut_config"]
        _add_itsar_heading(doc, cfg["heading"], 2)

        h = _add_itsar_subheading(doc, cfg["sub_4_1"]["heading"], 2)
        _keep_with_next(h)
        doc.add_paragraph(cfg["sub_4_1"]["body"])
        _add_bold_paragraph(doc, "Execution Command:")
        doc.add_paragraph(nmap_data["user_input"])
        _add_bold_paragraph(doc, "Executed Command Output:")
        doc.add_paragraph(nmap_data["terminal_output"] or "No output")

        if nmap_data.get("screenshot") and os.path.exists(nmap_data["screenshot"]):
            _add_screenshot_evidence_block(
                doc, "DUT Configuration : Nmap Scan Screenshot", nmap_data["screenshot"]
            )

        _add_itsar_subheading(doc, cfg["sub_4_2"]["heading"], 2)
        doc.add_paragraph(cfg["sub_4_2"]["body"])

        # Section 5 – Preconditions
        _add_itsar_heading(doc, s["5_preconditions"]["heading"], 2)
        for bullet in s["5_preconditions"]["bullets"]:
            doc.add_paragraph(f"• {bullet}")

        # Section 6 – Test Objective
        _add_itsar_heading(doc, s["6_test_objective"]["heading"], 2)
        doc.add_paragraph(s["6_test_objective"]["body"])

        # Section 7 – Test Scenario
        sc = s["7_test_scenario"]
        _add_itsar_heading(doc, sc["heading"], 2)

        _add_itsar_subheading(doc, sc["sub_7_1"]["heading"], 2)

        h = _add_itsar_subheading(doc, sc["sub_7_2"]["heading"], 2)
        _keep_with_next(h)
        p = doc.add_paragraph()
        _keep_with_next(p)
        testbed_img = sc["sub_7_2"]["testbed_image_path"]
        if os.path.exists(testbed_img):
            p.add_run().add_picture(testbed_img, width=Inches(6.5))

        _add_itsar_subheading(doc, sc["sub_7_3"]["heading"], 2)
        for bullet in sc["sub_7_3"]["bullets"]:
            doc.add_paragraph(f"• {bullet}")

        _add_itsar_subheading(doc, sc["sub_7_4"]["heading"], 2)
        for bullet in sc["sub_7_4"]["bullets"]:
            doc.add_paragraph(f"• {bullet}")

        # Section 8 – Expected Results
        _add_itsar_heading(doc, s["8_expected_results"]["heading"], 2)
        doc.add_paragraph(s["8_expected_results"]["body"])

    # ------------------------------------------------------------------
    # Section 9: SSH test execution
    # ------------------------------------------------------------------

    def _add_ssh_section(self, doc, r, tc1_ssh, tc2_ssh, tc3_ssh):
        s = self.content["sections"]["9_ssh_execution"]
        cipher_data     = r["cipher"]
        ssh_data        = r["ssh"]
        weak_cipher_result = r["weak_cipher"]

        _add_itsar_heading(doc, s["heading"], 2)

        # ---- TC-01: Cipher Support ----
        tc1 = s["tc1"]
        _add_itsar_subheading(doc, tc1["heading"], 2)
        _add_bold_paragraph(doc, tc1["name_label"])
        doc.add_paragraph(tc1["name"])
        _add_bold_paragraph(doc, tc1["desc_label"])
        doc.add_paragraph(tc1["desc"])
        _add_bold_paragraph(doc, tc1["steps_label"])

        p = doc.add_paragraph()
        p.add_run(tc1["steps_prefix"])
        bold_run = p.add_run(tc1["steps_command"])
        bold_run.bold = True
        p.add_run(tc1["steps_suffix"])
        for b in tc1["steps_bullets"]:
            doc.add_paragraph(f"• {b}")

        _add_bold_paragraph(doc, "Execution Command:")
        doc.add_paragraph(cipher_data["user_input"])
        _add_bold_paragraph(doc, "Executed Command Output:")
        doc.add_paragraph(cipher_data["terminal_output"] or "No output")

        details = cipher_data.get("details", {})
        sub = tc1["sub_sections"]

        h = _add_itsar_subheading(doc, sub["9_1_1"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(
            doc,
            details.get("encryption", {}).get("strong", []),
            details.get("encryption", {}).get("weak", []),
            "Strong Encryption", "Weak Encryption"
        )

        h = _add_itsar_subheading(doc, sub["9_1_2"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(
            doc,
            details.get("mac", {}).get("strong", []),
            details.get("mac", {}).get("weak", []),
            "Strong MAC", "Weak MAC"
        )

        h = _add_itsar_subheading(doc, sub["9_1_3"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(
            doc,
            details.get("kex", {}).get("strong", []),
            details.get("kex", {}).get("weak", []),
            "Strong KEX", "Weak KEX"
        )

        h = _add_itsar_subheading(doc, sub["9_1_4"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(
            doc,
            details.get("host_key", {}).get("strong", []),
            details.get("host_key", {}).get("weak", []),
            "Strong Host Key", "Weak Host Key"
        )
        _keep_with_next(doc.add_paragraph())
        doc.add_paragraph()

        if cipher_data.get("screenshot") and os.path.exists(cipher_data["screenshot"]):
            _add_screenshot_evidence_block(doc, tc1["screenshot_title"], cipher_data["screenshot"])

        # ---- TC-02: SSH data protection ----
        tc2 = s["tc2"]
        _add_itsar_subheading(doc, tc2["heading"], 2)
        _add_bold_paragraph(doc, tc2["name_label"])
        doc.add_paragraph(tc2["name"])
        _add_bold_paragraph(doc, tc2["desc_label"])
        doc.add_paragraph(tc2["desc"])
        _add_bold_paragraph(doc, tc2["steps_label"])
        for b in tc2["steps_bullets"]:
            doc.add_paragraph(f"• {b}")

        _add_bold_paragraph(doc, "Execution Command:")
        doc.add_paragraph(ssh_data["user_input"])

        h = _add_itsar_subheading(doc, tc2["sub_9_2_1"], 2); _keep_with_next(h)
        enc = doc.add_table(rows=4, cols=2)
        enc.style = "Table Grid"
        crypto = ssh_data.get("crypto_details", {})
        enc.cell(0, 0).text = "Protocol";              enc.cell(0, 1).text = crypto.get("protocol", "Not Found")
        enc.cell(1, 0).text = "Encryption Algorithm";  enc.cell(1, 1).text = crypto.get("cipher",   "Not Found")
        enc.cell(2, 0).text = "Key Exchange Algorithm"; enc.cell(2, 1).text = crypto.get("kex",       "Not Found")
        enc.cell(3, 0).text = "Host Key Algorithm";     enc.cell(3, 1).text = crypto.get("host_key",  "Not Found")
        _keep_with_next(doc.add_paragraph())

        # SSH screenshots
        if ssh_data.get("screenshots"):
            kex_algo    = crypto.get("kex",    "Unknown")
            cipher_algo = crypto.get("cipher", "Unknown")
            kex_status  = ssh_data["nist_validation"].get("kex",        "FAIL")
            enc_status  = ssh_data["nist_validation"].get("encryption",  "FAIL")
            kex_label   = "Secure" if kex_status == "PASS" else "Insecure"
            enc_label   = "Secure" if enc_status == "PASS" else "Insecure"

            titles = [
                tc2["screenshot_titles"][0],
                tc2["screenshot_titles"][1].format(kex_label=kex_label, kex_algo=kex_algo),
                tc2["screenshot_titles"][2].format(enc_label=enc_label, cipher_algo=cipher_algo),
            ]
            overviews = [
                tc2["screenshot_overviews"][0],
                tc2["screenshot_overviews"][1].format(kex_algo=kex_algo, kex_label=kex_label),
                tc2["screenshot_overviews"][2].format(cipher_algo=cipher_algo, enc_label=enc_label),
            ]

            for idx, img in enumerate(ssh_data["screenshots"]):
                if os.path.exists(img):
                    _add_screenshot_evidence_block(doc, titles[idx], img)
                    _add_itsar_heading(doc, f"9.2.{2 + idx}. Overview", 2)
                    doc.add_paragraph(overviews[idx])
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_after = Pt(18)

        # ---- Weak cipher (negative) tests ----
        weak_templates = s["weak_cipher"]["overview_templates"]

        if weak_cipher_result.get("screenshots"):
            results     = weak_cipher_result.get("results", [])
            screenshots = weak_cipher_result.get("screenshots", [])

            for idx, (res, img) in enumerate(zip(results, screenshots)):
                algo       = res.get("algorithm", "Unknown")
                algo_type  = res.get("type",      "Unknown")
                negotiated = res.get("negotiated", False)
                negotiation_text = "successfully negotiated" if negotiated else "rejected by the DUT"
                status_label = "Insecure"

                type_title_map = {
                    "cipher":   f"SSH Weak Cipher Attempt : {algo}",
                    "mac":      f"SSH Weak MAC Attempt : {algo}",
                    "kex":      f"SSH Weak Key Exchange Attempt : {algo}",
                    "host_key": f"SSH Weak Host Key Attempt : {algo}",
                }
                title = type_title_map.get(algo_type, f"SSH Weak Algorithm Attempt : {algo}")

                template = weak_templates.get(algo_type, weak_templates["default"])
                overview = template.format(
                    algo=algo,
                    negotiation_text=negotiation_text,
                    status_label=status_label,
                )

                if os.path.exists(img):
                    _add_screenshot_evidence_block(doc, title, img)
                    _add_itsar_heading(doc, f"9.2.{5 + idx}. Overview", 2)
                    doc.add_paragraph(overview)
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_after = Pt(18)

    # ------------------------------------------------------------------
    # Sections 10–11: SSH observation + result table
    # ------------------------------------------------------------------

    def _add_ssh_observation(self, doc, tc1, tc2, tc3):
        s = self.content["sections"]["10_ssh_observation"]
        labels = s["failed_test_labels"]
        failed = []
        if tc1 == "FAIL": failed.append(labels["tc1"])
        if tc2 == "FAIL": failed.append(labels["tc2"])
        if tc3 == "FAIL": failed.append(labels["tc3"])

        observation = (
            s["pass_text"] if not failed
            else s["fail_prefix"] + ", ".join(failed) + s["fail_suffix"]
        )
        _add_itsar_heading(doc, s["heading"], 2)
        doc.add_paragraph(observation)

    def _add_ssh_result_table(self, doc, tc1, tc2, tc3, weak_cipher_result):
        s = self.content["sections"]["11_ssh_results"]
        rows = s["rows"]
        weak_algos = [r["algorithm"] for r in weak_cipher_result.get("results", []) if r.get("negotiated")]

        h = _add_itsar_heading(doc, s["heading"], 2); _keep_with_next(h)

        rt = doc.add_table(rows=4, cols=4)
        rt.style = "Table Grid"
        _add_result_table_header(rt, s["table_headers"])

        # Row 1 – TC1
        rt.cell(1, 0).text = rows["tc1"]["sl"]
        rt.cell(1, 1).text = rows["tc1"]["name"]
        rt.cell(1, 2).text = tc1
        rt.cell(1, 3).text = rows["tc1"]["pass_remark"] if tc1 == "PASS" else rows["tc1"]["fail_remark"]

        # Row 2 – TC2
        rt.cell(2, 0).text = rows["tc2"]["sl"]
        rt.cell(2, 1).text = rows["tc2"]["name"]
        rt.cell(2, 2).text = tc2
        rt.cell(2, 3).text = rows["tc2"]["pass_remark"] if tc2 == "PASS" else rows["tc2"]["fail_remark"]

        # Row 3 – TC3
        rt.cell(3, 0).text = rows["tc3"]["sl"]
        rt.cell(3, 1).text = rows["tc3"]["name"]
        rt.cell(3, 2).text = tc3
        rt.cell(3, 3).text = (
            rows["tc3"]["pass_remark"] if tc3 == "PASS"
            else rows["tc3"]["fail_remark_prefix"] + ", ".join(weak_algos) + rows["tc3"]["fail_remark_suffix"]
        )

        _pad_data_rows(rt)
        _prevent_table_row_split(rt)

    # ------------------------------------------------------------------
    # Section 12: HTTPS test execution
    # ------------------------------------------------------------------

    def _add_https_section(self, doc, r, tc1_https, tc2_https):
        s = self.content["sections"]["12_https_execution"]
        https_cipher_data = r["https_cipher"]
        https_data        = r["https"]

        _add_itsar_heading(doc, s["heading"], 2)

        # ---- HTTPS TC-01: Cipher Support ----
        tc1 = s["tc1"]
        _add_itsar_subheading(doc, tc1["heading"], 2)
        _add_bold_paragraph(doc, tc1["name_label"])
        doc.add_paragraph(tc1["name"])
        _add_bold_paragraph(doc, tc1["desc_label"])
        doc.add_paragraph(tc1["desc"])
        _add_bold_paragraph(doc, tc1["steps_label"])

        p = doc.add_paragraph()
        p.add_run(tc1["steps_prefix"])
        p.add_run(tc1["steps_command"]).bold = True
        p.add_run(tc1["steps_suffix"])
        for b in tc1["steps_bullets"]:
            doc.add_paragraph(f"• {b}")

        _add_bold_paragraph(doc, "Execution Command:")
        doc.add_paragraph(https_cipher_data["user_input"])
        _add_bold_paragraph(doc, "Executed Command Output:")
        doc.add_paragraph(https_cipher_data["terminal_output"] or "No output")

        details = https_cipher_data.get("details", {})
        tls12   = details.get("TLSv1.2", {})
        tls13   = details.get("TLSv1.3", {})
        sub     = tc1["sub_sections"]

        h = _add_itsar_subheading(doc, sub["12_1_1"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(doc, tls12.get("encryption", {}).get("strong", []), tls12.get("encryption", {}).get("weak", []), "Strong Encryption", "Weak Encryption")

        h = _add_itsar_subheading(doc, sub["12_1_2"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(doc, tls12.get("mac", {}).get("strong", []), tls12.get("mac", {}).get("weak", []), "Strong MAC", "Weak MAC")

        h = _add_itsar_subheading(doc, sub["12_1_3"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(doc, tls12.get("kex", {}).get("strong", []), tls12.get("kex", {}).get("weak", []), "Strong KEX", "Weak KEX")

        h = _add_itsar_subheading(doc, sub["12_1_4"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(doc, tls13.get("encryption", {}).get("strong", []), tls13.get("encryption", {}).get("weak", []), "Strong Encryption", "Weak Encryption")

        h = _add_itsar_subheading(doc, sub["12_1_5"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(doc, tls13.get("mac", {}).get("strong", []), tls13.get("mac", {}).get("weak", []), "Strong MAC", "Weak MAC")

        h = _add_itsar_subheading(doc, sub["12_1_6"], 2); _keep_with_next(h)
        _add_two_col_cipher_table(doc, tls13.get("kex", {}).get("strong", []), tls13.get("kex", {}).get("weak", []), "Strong KEX", "Weak KEX")

        doc.add_paragraph()

        if https_cipher_data.get("screenshot") and os.path.exists(https_cipher_data["screenshot"]):
            _add_screenshot_evidence_block(doc, tc1["screenshot_title"], https_cipher_data["screenshot"])

        # ---- HTTPS TC-02: Data Protection ----
        tc2 = s["tc2"]
        _add_itsar_subheading(doc, tc2["heading"], 2)
        _add_bold_paragraph(doc, tc2["name_label"])
        doc.add_paragraph(tc2["name"])
        _add_bold_paragraph(doc, tc2["desc_label"])
        doc.add_paragraph(tc2["desc"])
        _add_bold_paragraph(doc, tc2["steps_label"])
        for b in tc2["steps_bullets"]:
            doc.add_paragraph(f"• {b}")

        _add_bold_paragraph(doc, "Execution Command:")
        doc.add_paragraph(https_data["user_input"])
        _add_bold_paragraph(doc, "Executed Command Output:")
        doc.add_paragraph(https_data["terminal_output"] or "No output")

        h = _add_itsar_subheading(doc, tc2["sub_12_2_1"], 2); _keep_with_next(h)
        tls_table = doc.add_table(rows=2, cols=2)
        tls_table.style = "Table Grid"
        crypto = https_data.get("crypto_details", {})
        tls_table.cell(0, 0).text = "Protocol";            tls_table.cell(0, 1).text = crypto.get("protocol", "Not Found")
        tls_table.cell(1, 0).text = "Encryption Algorithm"; tls_table.cell(1, 1).text = crypto.get("cipher",   "Not Found")
        _prevent_table_row_split(tls_table)
        _keep_with_next(doc.add_paragraph())

        # HTTPS screenshots
        if https_data.get("screenshots"):
            protocol      = crypto.get("protocol", "Unknown")
            cipher        = crypto.get("cipher",   "Unknown")
            proto_status  = https_data["nist_validation"].get("protocol", "FAIL")
            cipher_status = https_data["nist_validation"].get("cipher",   "FAIL")
            proto_label   = "Secure" if proto_status  == "PASS" else "Insecure"
            cipher_label  = "Secure" if cipher_status == "PASS" else "Insecure"

            titles = [
                tc2["screenshot_titles"][0],
                tc2["screenshot_titles"][1].format(cipher_label=cipher_label, cipher=cipher),
            ]
            overviews = [
                tc2["screenshot_overviews"][0],
                tc2["screenshot_overviews"][1].format(protocol=protocol, cipher=cipher, cipher_label=cipher_label),
            ]

            for idx, img in enumerate(https_data["screenshots"]):
                if os.path.exists(img):
                    _add_screenshot_evidence_block(doc, titles[idx], img)
                    _add_itsar_heading(doc, f"12.2.{2 + idx}. Overview", 2)
                    doc.add_paragraph(overviews[idx])
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_after = Pt(18)

    # ------------------------------------------------------------------
    # Sections 13–14: HTTPS observation + result table
    # ------------------------------------------------------------------

    def _add_https_observation(self, doc, tc1, tc2):
        s = self.content["sections"]["13_https_observation"]
        labels = s["failed_test_labels"]
        failed = []
        if tc1 == "FAIL": failed.append(labels["tc1"])
        if tc2 == "FAIL": failed.append(labels["tc2"])

        observation = (
            s["pass_text"] if not failed
            else s["fail_prefix"] + ", ".join(failed) + s["fail_suffix"]
        )
        _add_itsar_heading(doc, s["heading"], 2)
        doc.add_paragraph(observation)

    def _add_https_result_table(self, doc, tc1, tc2):
        s = self.content["sections"]["14_https_results"]
        rows = s["rows"]

        h = _add_itsar_heading(doc, s["heading"], 2); _keep_with_next(h)

        rt = doc.add_table(rows=3, cols=4)
        rt.style = "Table Grid"
        _add_result_table_header(rt, s["table_headers"])

        rt.cell(1, 0).text = rows["tc1"]["sl"]
        rt.cell(1, 1).text = rows["tc1"]["name"]
        rt.cell(1, 2).text = tc1
        rt.cell(1, 3).text = rows["tc1"]["pass_remark"] if tc1 == "PASS" else rows["tc1"]["fail_remark"]

        rt.cell(2, 0).text = rows["tc2"]["sl"]
        rt.cell(2, 1).text = rows["tc2"]["name"]
        rt.cell(2, 2).text = tc2
        rt.cell(2, 3).text = rows["tc2"]["pass_remark"] if tc2 == "PASS" else rows["tc2"]["fail_remark"]

        _pad_data_rows(rt)
        _prevent_table_row_split(rt)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _compute_weak_cipher_result(weak_cipher_result: dict) -> str:
        """Returns 'PASS' if no weak cipher was negotiated, else 'FAIL'."""
        for r in weak_cipher_result.get("results", []):
            if r.get("negotiated") is True:
                return "FAIL"
        return "PASS"