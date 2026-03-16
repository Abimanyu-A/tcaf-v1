"""
reporting/helpers.py
─────────────────────────────────────────────────────────────────────────────
All shared python-docx primitives for TCAF report generation.

Public API
──────────
Colours (RGBColor)
    PURPLE  LIGHT_PURPLE  ORANGE  DARK_GREY  MID_GREY
    TABLE_HEADER_BG  TABLE_ALT_BG  PASS_GREEN  FAIL_RED  WHITE

    Hex string variants (needed for XML shading attributes):
    HEX_PURPLE  HEX_PASS_GREEN  HEX_FAIL_RED  HEX_DARK_GREY  HEX_WHITE

Paragraph builders  (all append to doc in-place, return None)
    section_heading(doc, text)
    sub_heading(doc, text)
    tc_heading(doc, text)
    body_para(doc, text, *, bold, italic, color, font)
    label_value_para(doc, label, value, value_color)
    bullet_item(doc, text, bold)
    numbered_item(doc, text)
    spacer(doc, *, small, large)

Block builders
    terminal_block(doc, lines)          — monospace grey box
    add_screenshot(doc, path, width_inches)

Table builders
    two_col_info_table(doc, headers, col_widths, data_rows)
    four_col_header_table(doc, headers, data_rows, col_widths)
    status_result_table(doc, status, *, label, detail, wide)

Page builders
    build_front_page(doc, meta)
    build_doc_with_header_footer(dut_name, dut_version) → Document

Low-level helpers (also importable by clause modules)
    _style_cell(cell, fill_hex, border_hex, width_dxa, ...)
    _para_in_cell(cell, text, *, bold, italic, size_pt, color, center, font)
    _set_table_width(table, dxa)
    _set_col_widths(table, widths_dxa)
─────────────────────────────────────────────────────────────────────────────
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Colour palette ────────────────────────────────────────────────────────────
PURPLE          = RGBColor(0x4B, 0x00, 0x82)
ORANGE          = RGBColor(0x4B, 0x00, 0x82)
DARK_GREY       = RGBColor(0x40, 0x40, 0x40)
MID_GREY        = RGBColor(0x59, 0x59, 0x59)
PASS_GREEN      = RGBColor(0x00, 0x64, 0x00)
FAIL_RED        = RGBColor(0xCC, 0x00, 0x00)
WHITE           = RGBColor(0xFF, 0xFF, 0xFF)

# Hex strings — used as XML shading fill= attributes
HEX_PURPLE      = "4B0082"
HEX_PASS_GREEN  = "006400"
HEX_FAIL_RED    = "CC0000"
HEX_DARK_GREY   = "404040"
HEX_MID_GREY    = "595959"
HEX_WHITE       = "FFFFFF"

# Table colours (hex strings for shading fills)
TABLE_HEADER_BG = "4B0082"   # same as HEX_PURPLE — named alias kept for clarity
TABLE_ALT_BG    = "F5F0FB"
LIGHT_PURPLE    = "F3ECFA"
NOT_RUN_BG      = "F0F0F0"
NOT_RUN_COLOR   = RGBColor(0x77, 0x77, 0x77)


# ── Low-level XML cell helpers ────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color_hex: str = "CCCCCC"):
    tcPr  = cell._tc.get_or_add_tcPr()
    tcBrd = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:color"), color_hex)
        tcBrd.append(b)
    tcPr.append(tcBrd)


def _set_cell_width(cell, dxa: int):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _style_cell(cell, fill_hex: str, border_hex: str = "CCCCCC",
                width_dxa: int = None,
                top=80, bottom=80, left=120, right=120):
    """One-shot: shading + borders + optional width + margins."""
    _set_cell_shading(cell, fill_hex)
    _set_cell_borders(cell, border_hex)
    _set_cell_margins(cell, top, bottom, left, right)
    if width_dxa:
        _set_cell_width(cell, width_dxa)


def _para_in_cell(cell, text: str, *,
                  bold:    bool      = False,
                  italic:  bool      = False,
                  size_pt: float     = 9,
                  color:   RGBColor  = None,
                  center:  bool      = False,
                  font:    str       = "Arial"):
    """Replace the cell's first paragraph with a single styled run."""
    p = cell.paragraphs[0]
    p.clear()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold           = bold
    run.italic         = italic
    run.font.size      = Pt(size_pt)
    run.font.name      = font
    if color:
        run.font.color.rgb = color


def _set_table_width(table, dxa: int):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _set_col_widths(table, widths: list):
    """Inject w:tblGrid column definitions."""
    tbl     = table._tbl
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        tblGrid.append(col)
    tblPr = tbl.find(qn("w:tblPr"))
    idx   = list(tbl).index(tblPr) + 1 if tblPr is not None else 0
    tbl.insert(idx, tblGrid)


def _add_para_border_bottom(para, color_hex: str, size: int = 12):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Numbering (bullets + decimal) ─────────────────────────────────────────────
# We inject via raw XML into the document part because python-docx's high-level
# numbering API requires the Document to be constructed with numbering config
# upfront, but build_doc_with_header_footer() needs to return a plain Document.

_NUMBERING_XML = (
    '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:abstractNum w:abstractNumId="0">'
      '<w:lvl w:ilvl="0">'
        '<w:start w:val="1"/>'
        '<w:numFmt w:val="bullet"/>'
        '<w:lvlText w:val="&#x2022;"/>'
        '<w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
      '</w:lvl>'
    '</w:abstractNum>'
    '<w:abstractNum w:abstractNumId="1">'
      '<w:lvl w:ilvl="0">'
        '<w:start w:val="1"/>'
        '<w:numFmt w:val="decimal"/>'
        '<w:lvlText w:val="%1."/>'
        '<w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
      '</w:lvl>'
    '</w:abstractNum>'
    '<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
    '<w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num>'
    '</w:numbering>'
)

def _ensure_numbering(doc: Document):
    """Inject bullet + decimal numbering into the document part if absent."""
    from lxml import etree
    part = doc.part
    if part.numbering_part is not None:
        return
    from docx.opc.part import Part
    numbering_part = Part(
        "/word/numbering.xml",
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.numbering+xml",
        etree.tostring(etree.fromstring(_NUMBERING_XML.encode())),
        doc.part.package,
    )
    doc.part.relate_to(
        numbering_part,
        "http://schemas.openxmlformats.org/officeDocument/"
        "2006/relationships/numbering",
    )


def _apply_list_style(para, num_id: int):
    """Attach numId (1=bullet, 2=decimal) to a paragraph."""
    pPr   = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl  = OxmlElement("w:ilvl");  ilvl.set(qn("w:val"), "0")
    nid   = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id))
    numPr.append(ilvl); numPr.append(nid)
    pPr.append(numPr)


# ── Paragraph builders ────────────────────────────────────────────────────────

def section_heading(doc: Document, text: str):
    """Purple bold heading with a purple underline border."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(14)
    run.font.name = "Arial"; run.font.color.rgb = PURPLE
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after  = Pt(6)
    _add_para_border_bottom(p, HEX_PURPLE, size=12)


def sub_heading(doc: Document, text: str):
    """Orange bold sub-heading."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(12)
    run.font.name = "Arial"; run.font.color.rgb = ORANGE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)


def tc_heading(doc: Document, text: str):
    """Purple bold test-case heading."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(11)
    run.font.name = "Arial"; run.font.color.rgb = PURPLE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)


def body_para(doc: Document, text: str, *,
              bold:   bool     = False,
              italic: bool     = False,
              color:  RGBColor = None,
              font:   str      = "Arial"):
    """Standard body paragraph. Defaults to DARK_GREY."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(10); run.font.name = font
    run.font.color.rgb = color if color is not None else DARK_GREY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)


def label_value_para(doc: Document, label: str, value: str,
                     value_color: RGBColor = None):
    """Bold label + plain value on the same line."""
    p     = doc.add_paragraph()
    r_lbl = p.add_run(f"{label}: ")
    r_lbl.bold = True; r_lbl.font.size = Pt(10)
    r_lbl.font.name = "Arial"; r_lbl.font.color.rgb = MID_GREY

    r_val = p.add_run(value or "N/A")
    r_val.font.size = Pt(10); r_val.font.name = "Arial"
    r_val.font.color.rgb = value_color if value_color else DARK_GREY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)


def bullet_item(doc: Document, text: str, bold: bool = False):
    """Bullet-list paragraph."""
    _ensure_numbering(doc)
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(10)
    run.font.name = "Arial"; run.font.color.rgb = DARK_GREY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    _apply_list_style(p, num_id=1)


def numbered_item(doc: Document, text: str):
    """Decimal-list paragraph."""
    _ensure_numbering(doc)
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Arial"
    run.font.color.rgb = DARK_GREY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    _apply_list_style(p, num_id=2)


def spacer(doc: Document, *, small: bool = False, large: bool = False):
    """Empty paragraph for vertical whitespace."""
    p = doc.add_paragraph()
    p.add_run("")
    if small:
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
    elif large:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(10)
    else:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)


# ── Block builders ────────────────────────────────────────────────────────────

def terminal_block(doc: Document, lines: list):
    """
    Monospace grey box for command output.

    `lines` should already be cleaned (stripped, capped at 40).
    Each element becomes its own paragraph inside the cell.
    """
    table = doc.add_table(rows=1, cols=1)
    _set_table_width(table, 9360)
    _set_col_widths(table, [9360])

    cell = table.rows[0].cells[0]
    _style_cell(cell, "EEEEEE", "AAAAAA", 9360,
                top=100, bottom=100, left=180, right=180)

    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)


def add_screenshot(doc: Document, path: str, width_inches: float = 5.5):
    """
    Insert an image, centred. Silently skips if path is None or missing.

    The runner embeds Python object reprs in screenshot paths, e.g.:
        output/runs/.../1.1.1/<Module.ClassName object at 0x7f...>/screenshots/file.png

    Those paths literally cannot exist. This function resolves them via
    _resolve_screenshot_path() and silently skips if the file isn't found,
    so the report generates cleanly even with broken paths.
    """
    clean = _resolve_screenshot_path(path)
    if not clean:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(clean, width=Inches(width_inches))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def _resolve_screenshot_path(raw: str) -> str | None:
    """
    The runner produces paths like:
        output/runs/2026-03-13_10-10-42-1.1.1/1.1.1/
        <clauses.clause_1_1_1.tc1_ssh_first_connection.TC1SSHFirstConnection
         object at 0x7fefd9311fd0>/screenshots/packet_frame_23.png

    The <...object at 0x...> segment is a Python repr — it is NOT a real
    directory name. We strip everything between < > and rebuild the path
    from the segments before and after it, then check if the result exists.

    Falls back to checking raw and cwd-joined raw before giving up.
    Returns None (never raises) so the report always generates cleanly.
    """
    if not raw:
        return None

    # Fast path: already valid
    if os.path.exists(raw):
        return raw

    # Try stripping the <...object at 0x...> segment
    import re
    cleaned = re.sub(r"<[^>]+>", "", raw)       # remove < ... >
    cleaned = re.sub(r"/+", "/", cleaned)        # collapse double slashes
    cleaned = cleaned.strip("/")

    if os.path.exists(cleaned):
        return cleaned

    joined = os.path.join(os.getcwd(), cleaned)
    if os.path.exists(joined):
        return joined

    return None


# ── Table builders ────────────────────────────────────────────────────────────

def two_col_info_table(doc: Document,
                       headers: list,
                       col_widths: list,
                       data_rows: list):
    """
    Two-column table: purple header row, shaded first column in data rows.

    headers     2 strings
    col_widths  2 DXA ints  (should sum to ≤ 9360)
    data_rows   list of (col0, col1) tuples
    """
    table = doc.add_table(rows=0, cols=2)
    _set_table_width(table, sum(col_widths))
    _set_col_widths(table, col_widths)

    # Header
    hdr = table.add_row()
    for ci, (h, w) in enumerate(zip(headers, col_widths)):
        c = hdr.cells[ci]
        _style_cell(c, TABLE_HEADER_BG, HEX_PURPLE, w)
        _para_in_cell(c, h, bold=True, color=WHITE, center=True)

    # Data
    for ri, row_vals in enumerate(data_rows):
        row = table.add_row()
        for ci, (val, w) in enumerate(zip(row_vals, col_widths)):
            c    = row.cells[ci]
            fill = LIGHT_PURPLE if ci == 0 else ("FFFFFF" if ri % 2 == 0 else "FAFAFA")
            _style_cell(c, fill, "CCCCCC", w)
            _para_in_cell(c, str(val or "N/A"), bold=(ci == 0), color=DARK_GREY)


def four_col_header_table(doc: Document,
                          headers: list,
                          data_rows: list,
                          col_widths: list = None):
    """
    Four-column table — purple header, alternating body rows.
    Default col_widths: equal quarters of 9360 DXA.
    """
    if col_widths is None:
        col_widths = [2340, 2340, 2340, 2340]

    table = doc.add_table(rows=0, cols=4)
    _set_table_width(table, sum(col_widths))
    _set_col_widths(table, col_widths)

    hdr = table.add_row()
    for ci, (h, w) in enumerate(zip(headers, col_widths)):
        c = hdr.cells[ci]
        _style_cell(c, TABLE_HEADER_BG, HEX_PURPLE, w)
        _para_in_cell(c, h, bold=True, color=WHITE, center=True)

    for ri, row_vals in enumerate(data_rows):
        row  = table.add_row()
        fill = "FFFFFF" if ri % 2 == 0 else TABLE_ALT_BG
        for ci, (val, w) in enumerate(zip(row_vals, col_widths)):
            c = row.cells[ci]
            _style_cell(c, fill, "CCCCCC", w)
            _para_in_cell(c, str(val or ""), color=DARK_GREY, center=(ci == 0))


def status_result_table(doc: Document,
                        status: str, *,
                        label:  str  = "Test Result",
                        detail: str  = "",
                        wide:   bool = False):
    """
    PASS/FAIL/NOT RUN badge as a coloured 2-column table.

    wide=False  inline badge  (label 3800 + status 1200 DXA)
    wide=True   full-width banner  (label+status 5000 + detail 4360 DXA)
    """
    sk = status.upper()
    if sk == "PASS":
        color_hex = HEX_PASS_GREEN
        color_rgb = PASS_GREEN
        bg_hex    = "E8F5E9"
    elif sk == "NOT RUN":
        color_hex = "777777"
        color_rgb = NOT_RUN_COLOR
        bg_hex    = NOT_RUN_BG
    else:   # FAIL
        color_hex = HEX_FAIL_RED
        color_rgb = FAIL_RED
        bg_hex    = "FFEBEE"

    if wide:
        col_widths  = [5000, 4360]
        left_text   = f"{label}:  {sk}"
        right_text  = detail
        left_bold   = True
        right_bold  = False
        left_size   = 12
        right_size  = 10
    else:
        col_widths  = [3800, 1200]
        left_text   = label
        right_text  = sk
        left_bold   = True
        right_bold  = True
        left_size   = 10
        right_size  = 10

    table = doc.add_table(rows=1, cols=2)
    _set_table_width(table, sum(col_widths))
    _set_col_widths(table, col_widths)

    row = table.rows[0]
    for ci, (text, bold, size, w) in enumerate(zip(
        [left_text,  right_text],
        [left_bold,  right_bold],
        [left_size,  right_size],
        col_widths,
    )):
        c = row.cells[ci]
        _style_cell(c, bg_hex, color_hex, w,
                    top=100, bottom=100, left=160, right=160)
        _para_in_cell(c, text, bold=bold, color=color_rgb,
                      size_pt=size, center=(not wide and ci == 1))


# ── Page-level builders ───────────────────────────────────────────────────────

def build_front_page(doc: Document, meta: dict):
    """
    Renders the front page content into doc.

    meta keys: dut_name, start_time, end_time
    (final_result, os_hash etc. come from the preface section, not here)
    """
    dut_name   = meta.get("dut_name",   "Device Under Test")
    start_time = meta.get("start_time", "N/A")
    end_time   = meta.get("end_time",   "N/A")

    spacer(doc, large=True)
    spacer(doc, large=True)

    def _centered(text, size_pt, color, bold=False):
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold; run.font.size = Pt(size_pt)
        run.font.name = "Arial"; run.font.color.rgb = color
        return p

    _centered("TSTL Evaluation Test Report", 26, PURPLE, bold=True)
    _centered("for", 18, MID_GREY)
    p = _centered(dut_name, 24, PURPLE, bold=True)
    p.paragraph_format.space_after = Pt(14)
    _add_para_border_bottom(p, HEX_PURPLE, size=12)

    spacer(doc, large=True)


    # Prepared for
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared For:  "); r.bold = True
    r.font.size = Pt(12); r.font.name = "Arial"; r.font.color.rgb = MID_GREY
    r2 = p.add_run(
        "National Centre of Communications Security, Bengaluru, "
        "Department of Telecom, Ministry of Communications, Government of India"
    )
    r2.font.size = Pt(10); r2.font.name = "Arial"; r2.font.color.rgb = MID_GREY

    spacer(doc, large=True)

    four_col_header_table(
        doc,
        headers    = ["Document No.", "Created By", "Reviewed By", "Approved By"],
        data_rows  = [("/IP ROUTER/1.1", "Tester", "Reviewer", "Approver")],
        col_widths = [2340, 2340, 2340, 2340],
    )
    spacer(doc, small=True)
    four_col_header_table(
        doc,
        headers    = ["Test conducted by", "Test conducted on",
                      "Test reviewed by",  "Test reviewed on"],
        data_rows  = [("Tester", start_time, "Reviewer", end_time)],
        col_widths = [2340, 2340, 2340, 2340],
    )


def build_doc_with_header_footer(dut_name: str, dut_version: str) -> Document:
    """
    Return a new Document pre-configured with:
      • US Letter, 0.75-inch margins
      • Purple-themed header: copyright left, clause right
      • Page-numbered footer: DUT info left, Page N of M right
    """
    doc     = Document()
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

    # ── Header ────────────────────────────────────────────────────────────
    header = section.header
    hp     = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    _add_para_border_bottom(hp, HEX_PURPLE, size=6)
    hp.paragraph_format.space_after = Pt(4)

    # right-align tab stop at content width
    pPr  = hp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), "9360")
    tabs.append(tab); pPr.append(tabs)

    r2 = hp.add_run("\tReport - ITSAR 1.1.1")
    r2.bold = True; r2.font.size = Pt(8)
    r2.font.name = "Arial"; r2.font.color.rgb = PURPLE

    # ── Footer ────────────────────────────────────────────────────────────
    footer = section.footer
    fp     = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.paragraph_format.space_before = Pt(4)

    # top border
    pPr  = fp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "6")
    top.set(qn("w:color"), HEX_PURPLE)
    pBdr.append(top); pPr.append(pBdr)

    # right-align tab
    tabs2 = OxmlElement("w:tabs")
    tab2  = OxmlElement("w:tab")
    tab2.set(qn("w:val"), "right"); tab2.set(qn("w:pos"), "9360")
    tabs2.append(tab2); pPr.append(tabs2)

    r3 = fp.add_run(f"DUT: {dut_name}  |  Firmware: {dut_version}")
    r3.font.size = Pt(8); r3.font.name = "Arial"; r3.font.color.rgb = MID_GREY
    fp.add_run("\t").font.size = Pt(8)

    # Page N of M via field chars
    p_el = fp._p

    def _fld(type_):
        r  = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), type_)
        r.append(fc); return r

    def _instr(text):
        r = OxmlElement("w:r")
        i = OxmlElement("w:instrText"); i.text = text
        i.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(i); return r

    def _txt(text):
        r  = OxmlElement("w:r")
        rp = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16")
        rp.append(sz); r.append(rp)
        t  = OxmlElement("w:t"); t.text = text; r.append(t); return r

    for node in [
        _txt("Page "),
        _fld("begin"), _instr(" PAGE "), _fld("separate"), _txt("1"), _fld("end"),
        _txt(" of "),
        _fld("begin"), _instr(" NUMPAGES "), _fld("separate"), _txt("1"), _fld("end"),
    ]:
        p_el.append(node)

    return doc