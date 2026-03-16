from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_itsar_subheading(doc, text, level):
    para = doc.add_heading(text, level=level)
    run = para.runs[0]
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.color.rgb = RGBColor(75, 0, 130)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(8)
    return para

def add_itsar_heading(doc, text, level):
    para = doc.add_paragraph()

    # Font size based on level (matches your old logic)
    font_size = Pt(16 if level == 1 else 14)

    run = para.add_run(text)
    run.bold = True
    run.font.size = font_size
    run.font.color.rgb = RGBColor(75, 0, 130)  # same blue you used

    # Spacing (tight like ITSAR image)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(2)

    # Add bottom border (horizontal line)
    p = para._p
    pPr = p.get_or_add_pPr()

    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')        # line thickness
    bottom.set(qn('w:space'), '2')      # gap between text and line
    bottom.set(qn('w:color'), '4B0082') # same heading blue

    pBdr.append(bottom)
    pPr.append(pBdr)

    return para