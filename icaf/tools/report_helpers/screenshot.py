from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from .tables import prevent_table_row_split


def add_screenshot_evidence_block(doc, title, image_path):
    TABLE_WIDTH = Inches(7.8)
    IMAGE_WIDTH = Inches(6.2)

    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    # PREVENT TABLE FROM BREAKING ACROSS PAGES
    prevent_table_row_split(table)

    # Lock table width
    table.columns[0].width = TABLE_WIDTH
    for row in table.rows:
        row.cells[0].width = TABLE_WIDTH

    # ---------- COMMON CELL STYLE ----------
    for row in table.rows:
        cell = row.cells[0]
        tcPr = cell._tc.get_or_add_tcPr()

        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F3ECFA')
        tcPr.append(shd)

        cell.top_margin = Inches(0.2)
        cell.bottom_margin = Inches(0.2)
        cell.left_margin = Inches(0.3)
        cell.right_margin = Inches(0.3)

    # ---------- TITLE ----------
    title_cell = table.cell(0, 0)
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True

    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 0, 130)

    # ---------- IMAGE ----------
    img_cell = table.cell(1, 0)
    p_img = img_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.keep_together = True
    p_img.add_run().add_picture(image_path, width=IMAGE_WIDTH)

    # ---------- PURPLE BORDER ----------
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), '4B0082')
        tblBorders.append(border)

    tblPr.append(tblBorders)

    return table